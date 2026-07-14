"""
Fayllarni qayta ishlash logikasi. Celery task'lar shu funksiyalarni chaqiradi.

MUHIM: bu fayl R2 (masofaviy) storage bilan ishlaydigan qilib yozilgan.
`job.input_file.path` kabi "lokal disk yo'li" endi ishlatilmaydi, chunki
R2'da fayl fizik diskda emas -- faqat network orqali o'qish/yozish mumkin.
Shuning uchun har bir job uchun:
  1. input faylni vaqtinchalik lokal nusxaga ko'chiramiz (kutubxonalar
     ko'pincha haqiqiy fayl yo'lini talab qiladi),
  2. natijani ham avval vaqtinchalik lokal faylga yozamiz,
  3. tayyor bo'lgach, uni Django Storage API orqali (`job.result_file.save`)
     to'g'ri joyga (R2'ga yoki lokal `media/`ga -- sozlamaga qarab)
     yuklaymiz,
  4. vaqtinchalik fayllarni o'chirib tashlaymiz.
"""

import os
import shutil
import tempfile
from contextlib import contextmanager

import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from django.core.files import File


@contextmanager
def _local_input_copy(job):
    """
    job.input_file'ni (R2'da yoki lokal diskda bo'lishidan qat'iy nazar)
    vaqtinchalik lokal faylga ko'chirib beradi va shu faylning yo'lini
    qaytaradi. `with` blokidan chiqilgach, vaqtinchalik fayl avtomatik
    o'chiriladi.
    """
    suffix = os.path.splitext(job.input_file.name)[1]
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(fd, "wb") as tmp_file, job.input_file.open("rb") as src:
            shutil.copyfileobj(src, tmp_file)
        yield tmp_path
    finally:
        os.remove(tmp_path)


@contextmanager
def _local_output_path(filename):
    """
    Natija faylini vaqtinchalik yozish uchun bo'sh papka+fayl yo'lini
    beradi. Chaqiruvchi shu yo'lga fayl yozadi, keyin `_save_result`
    orqali storage'ga yuklaydi. `with` blokidan chiqilgach, vaqtinchalik
    papka butunlay o'chiriladi.
    """
    tmp_dir = tempfile.mkdtemp()
    try:
        yield os.path.join(tmp_dir, filename)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _save_result(job, filename, local_path):
    """
    Tayyor natija faylini job.result_file maydoniga saqlaydi.
    `.save()` -- Django Storage API'sining o'zi, shuning uchun fayl
    sozlamaga qarab avtomatik to'g'ri joyga (R2'ga yoki lokal media/ga)
    yuklanadi. `save=False` beryapmiz, chunki job'ni (status bilan
    birga) bitta marta tasks.py o'zi saqlaydi -- ortiqcha DB so'rovi
    bo'lmasin deb.
    """
    with open(local_path, "rb") as f:
        job.result_file.save(filename, File(f), save=False)


def process_excel_clean(job):
    """
    Excel faylni o'qiydi: bo'sh qatorlarni va to'liq takrorlangan qatorlarni
    o'chiradi, natijani yangi xlsx sifatida saqlaydi.
    options: {"dedup_columns": ["col1", "col2"]} -- ixtiyoriy, berilmasa barcha
    ustunlar bo'yicha dedup qilinadi.
    """
    with _local_input_copy(job) as input_path, _local_output_path("cleaned.xlsx") as out_path:
        df = pd.read_excel(input_path)
        df = df.dropna(how="all")

        dedup_cols = job.options.get("dedup_columns") if job.options else None
        if dedup_cols:
            df = df.drop_duplicates(subset=dedup_cols)
        else:
            df = df.drop_duplicates()

        df.to_excel(out_path, index=False)
        _save_result(job, "cleaned.xlsx", out_path)


def process_excel_to_pdf(job):
    """
    Excel faylni o'qiydi va oddiy jadval ko'rinishidagi PDF hisobot yaratadi.
    """
    with _local_input_copy(job) as input_path, _local_output_path("report.pdf") as out_path:
        df = pd.read_excel(input_path)

        doc = SimpleDocTemplate(out_path, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = [Paragraph("FileFlow hisobot", styles["Title"]), Spacer(1, 12)]

        data = [list(df.columns)] + df.astype(str).values.tolist()

        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d3748")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7fafc")]),
        ]))
        elements.append(table)

        doc.build(elements)
        _save_result(job, "report.pdf", out_path)


def process_pdf_table_extract(job):
    """
    PDF ichidagi jadval(lar)ni topib, har birini bitta Excel faylga
    alohida sheet sifatida yozadi.
    """
    with _local_input_copy(job) as input_path, _local_output_path("extracted_tables.xlsx") as out_path:
        wb = Workbook()
        wb.remove(wb.active)  # default bo'sh sheet'ni o'chiramiz

        sheet_count = 0
        with pdfplumber.open(input_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables, start=1):
                    sheet_count += 1
                    ws = wb.create_sheet(title=f"p{page_num}_t{t_idx}"[:31])
                    for row in table:
                        ws.append([cell if cell is not None else "" for cell in row])

        if sheet_count == 0:
            ws = wb.create_sheet(title="empty")
            ws.append(["PDF ichida jadval topilmadi"])

        wb.save(out_path)
        _save_result(job, "extracted_tables.xlsx", out_path)


def process_docx_to_pdf(job):
    """
    Word (.docx) hujjatni o'qiydi va oddiy PDF hisobotga aylantiradi.
    Faqat paragraf matnlarini va jadval hujayralarini ko'chiradi -- murakkab
    formatlash (rasm, ustunlar, stil) saqlanmaydi, faqat matn mazmuni ko'chadi.
    """
    from xml.sax.saxutils import escape

    with _local_input_copy(job) as input_path, _local_output_path("converted.pdf") as out_path:
        document = Document(input_path)

        doc = SimpleDocTemplate(out_path, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                elements.append(Paragraph(escape(text), styles["Normal"]))
                elements.append(Spacer(1, 8))

        for table in document.tables:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            if not data:
                continue
            pdf_table = Table(data, repeatRows=1)
            pdf_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d3748")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(pdf_table)
            elements.append(Spacer(1, 12))

        if not elements:
            elements.append(Paragraph("(Hujjat bo'sh)", styles["Normal"]))

        doc.build(elements)
        _save_result(job, "converted.pdf", out_path)


def process_pdf_text_extract(job):
    """
    PDF ichidagi barcha sahifalar matnini ketma-ket o'qib, Word (.docx)
    hujjatga yozadi. Jadval izlamaydi (buning uchun process_pdf_table_extract
    bor) -- bu funksiya "oddiy" matnli PDF'lar uchun.
    """
    with _local_input_copy(job) as input_path, _local_output_path("extracted_text.docx") as out_path:
        document = Document()

        found_any_text = False
        with pdfplumber.open(input_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = text.strip()
                if not text:
                    continue
                found_any_text = True
                document.add_heading(f"{page_num}-sahifa", level=2)
                for line in text.split("\n"):
                    if line.strip():
                        document.add_paragraph(line.strip())

        if not found_any_text:
            document.add_paragraph("PDF ichida matn topilmadi.")

        document.save(out_path)
        _save_result(job, "extracted_text.docx", out_path)


PROCESSORS = {
    "excel_clean": process_excel_clean,
    "excel_to_pdf": process_excel_to_pdf,
    "pdf_table_extract": process_pdf_table_extract,
    "docx_to_pdf": process_docx_to_pdf,
    "pdf_text_extract": process_pdf_text_extract,
}